import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook


BASE_DIR = Path(__file__).resolve().parent.parent.parent
PROCESSED_DIR = BASE_DIR / "data" / "processed"
CHUNKS_FILE = PROCESSED_DIR / "chunks.json"
EXPORT_DIR = BASE_DIR / "data" / "exports"


def ensure_export_dir():
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def load_chunks(chunks_file: Path = CHUNKS_FILE):
    if not chunks_file.exists():
        return []
    with chunks_file.open("r", encoding="utf-8") as f:
        return json.load(f)


def export_chunks_to_docx(chunks, out_path: Path | str = None):
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "chunks_export.docx"
    doc = Document()
    doc.add_heading("Document Chunks Export", level=1)
    for c in chunks:
        doc.add_heading(f"{c.get('source')} — {c.get('id')}", level=3)
        doc.add_paragraph(c.get('text', ''))
    doc.save(out_path)
    return out_path


def export_chunks_to_excel(chunks, out_path: Path | str = None):
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "chunks_export.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "chunks"
    ws.append(["id", "text", "source"])
    for c in chunks:
        ws.append([c.get('id'), c.get('text'), c.get('source')])
    wb.save(out_path)
    return out_path


def export_qa_to_docx(qa_rows, out_path: Path | str = None):
    """qa_rows: list of dicts with keys: question, answer, source"""
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "qa_export.docx"
    doc = Document()
    doc.add_heading("QA Export", level=1)
    for r in qa_rows:
        doc.add_heading(r.get('question', ''), level=3)
        doc.add_paragraph(r.get('answer', ''))
        if r.get('source'):
            doc.add_paragraph(f"Source: {r.get('source')}")
    doc.save(out_path)
    return out_path


def export_qa_to_excel(qa_rows, out_path: Path | str = None):
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "qa_export.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "qa"
    ws.append(["question", "answer", "source"])
    for r in qa_rows:
        ws.append([r.get('question'), r.get('answer'), r.get('source')])
    wb.save(out_path)
    return out_path


if __name__ == "__main__":
    chunks = load_chunks()
    if not chunks:
        print("No chunks found to export.")
    else:
        docx_path = export_chunks_to_docx(chunks)
        xlsx_path = export_chunks_to_excel(chunks)
        print(f"Exported {len(chunks)} chunks to:\n - {docx_path}\n - {xlsx_path}")
