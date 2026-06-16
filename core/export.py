import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook

BASE_DIR = Path.cwd()
EXPORT_DIR = BASE_DIR / "data" / "exports"
CHUNKS_FILE = BASE_DIR / "data" / "processed" / "chunks.json"


def ensure_export_dir() -> None:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def load_chunks() -> list[dict]:
    if not CHUNKS_FILE.exists():
        return []
    with CHUNKS_FILE.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def export_chunks_to_docx(chunks: list[dict], out_path: Path | str = None) -> Path:
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "chunks_export.docx"
    doc = Document()
    doc.add_heading("Knowledge Base Chunk Export", level=1)
    for chunk in chunks:
        doc.add_heading(f"{chunk.get('source')} ({chunk.get('id')})", level=3)
        doc.add_paragraph(chunk.get('text', ''))
    doc.save(out_path)
    return out_path


def export_chunks_to_excel(chunks: list[dict], out_path: Path | str = None) -> Path:
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "chunks_export.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "chunks"
    ws.append(["id", "source", "text"])
    for chunk in chunks:
        ws.append([chunk.get('id'), chunk.get('source'), chunk.get('text')])
    wb.save(out_path)
    return out_path


def export_qa_to_docx(rows: list[dict], out_path: Path | str = None) -> Path:
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "qa_export.docx"
    doc = Document()
    doc.add_heading("QA Export", level=1)
    for row in rows:
        doc.add_heading(row.get('question', ''), level=3)
        doc.add_paragraph(row.get('answer', ''))
        if row.get('source'):
            doc.add_paragraph(f"Source: {row.get('source')}")
    doc.save(out_path)
    return out_path


def export_qa_to_excel(rows: list[dict], out_path: Path | str = None) -> Path:
    ensure_export_dir()
    out_path = Path(out_path) if out_path else EXPORT_DIR / "qa_export.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "qa"
    ws.append(["question", "answer", "source"])
    for row in rows:
        ws.append([row.get('question'), row.get('answer'), row.get('source')])
    wb.save(out_path)
    return out_path
