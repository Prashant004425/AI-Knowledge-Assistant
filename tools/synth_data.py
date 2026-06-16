import json
import os
from datetime import date

import fitz
import pandas as pd
from docx import Document


BASE_DATE = date.today().isoformat()
RAW_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", "data", "raw"))


def ensure_raw_dir():
    os.makedirs(RAW_DIR, exist_ok=True)


def write_markdown(filename: str, title: str, tags: list[str], body: str):
    metadata = {
        "title": title,
        "tags": tags,
        "date": BASE_DATE,
    }
    content = f"""```json
{json.dumps(metadata, indent=2)}
```

# {title}

{body}
"""
    path = os.path.join(RAW_DIR, filename)
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(content)
    return path


def write_pdf(filename: str, title: str, tags: list[str], body: str):
    metadata = {
        "title": title,
        "tags": tags,
        "date": BASE_DATE,
    }
    doc = fitz.open()
    page = doc.new_page()
    text = f"METADATA:\n{json.dumps(metadata, indent=2)}\n\n{title}\n\n{body}"
    page.insert_text(fitz.Point(50, 50), text, fontsize=11)
    path = os.path.join(RAW_DIR, filename)
    doc.save(path)
    doc.close()
    return path


def write_docx(filename: str, title: str, tags: list[str], body: str):
    metadata = {
        "title": title,
        "tags": tags,
        "date": BASE_DATE,
    }
    document = Document()
    document.add_paragraph("METADATA:")
    document.add_paragraph(json.dumps(metadata, indent=2))
    document.add_heading(title, level=1)
    for paragraph in body.split("\n\n"):
        document.add_paragraph(paragraph)
    path = os.path.join(RAW_DIR, filename)
    document.save(path)
    return path


def write_csv(filename: str, title: str, tags: list[str], rows: list[dict[str, str]]):
    metadata = {
        "title": title,
        "tags": tags,
        "date": BASE_DATE,
    }
    path = os.path.join(RAW_DIR, filename)
    with open(path, "w", encoding="utf-8", newline="") as handle:
        handle.write(f"# {json.dumps(metadata)}\n")
    df = pd.DataFrame(rows)
    df.to_csv(path, mode="a", index=False)
    return path


def generate_sample_documents():
    ensure_raw_dir()

    docs = [
        {
            "writer": write_markdown,
            "filename": "onboarding_welcome.md",
            "title": "Onboarding Guide",
            "tags": ["onboarding", "hr", "welcome"],
            "body": (
                "Welcome to the company! This onboarding document explains the first-week "
                "checklist, account setup steps, and team introductions.\n\n"
                "New hires should complete their profile, request access to internal tools, "
                "and join the onboarding Slack channel."
            ),
        },
        {
            "writer": write_markdown,
            "filename": "flocard_api_guide.md",
            "title": "FloCard API Guide",
            "tags": ["payments", "api"],
            "body": (
                "FloCard is the internal payment orchestration platform.\n\n"
                "Authentication uses OAuth2 bearer tokens. Use the `/v1/transactions` endpoint "
                "for payment capture and `/v1/refunds` for refund requests."
            ),
        },
        {
            "writer": write_pdf,
            "filename": "company_policy_code_of_conduct.pdf",
            "title": "Code of Conduct",
            "tags": ["policy", "compliance"],
            "body": (
                "This document defines acceptable behavior in the workplace. Employees should "
                "treat each other with respect, report issues early, and follow the security policy."
            ),
        },
        {
            "writer": write_docx,
            "filename": "faq_employee_tools.docx",
            "title": "Employee Tools FAQ",
            "tags": ["faq", "tools", "support"],
            "body": (
                "Q: How do I request a new laptop?\nA: Submit a ticket through IT support.\n\n"
                "Q: Where do I find the VPN configuration?\nA: See the network access page in the intranet."
            ),
        },
        {
            "writer": write_csv,
            "filename": "sop_content_review.csv",
            "title": "Content Review SOP",
            "tags": ["sop", "content", "workflow"],
            "rows": [
                {"step": "Draft document", "owner": "Author", "deadline": "2 days"},
                {"step": "Review draft", "owner": "Editor", "deadline": "1 day"},
                {"step": "Publish", "owner": "Ops", "deadline": "Same day"},
            ],
        },        
    ]

    paths = []
    for doc in docs:
        kwargs = {k: doc[k] for k in doc if k not in {"writer"}}
        paths.append(doc["writer"](**kwargs))

    return paths


if __name__ == "__main__":
    generated = generate_sample_documents()
    print("Generated files:")
    for path in generated:
        print(path)
